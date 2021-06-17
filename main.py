from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import re
import pandas as pd
import os
from tqdm import tqdm
from colorama import Fore
import collections

from win32com.client import Dispatch


# def respec(match):
#     return "{}_hhh".format(match.group(0))
#
# def delete_paragraph(paragraph):
#     p = paragraph._element
#     p.getparent().remove(p)
#     # p._p = p._element = None
#     paragraph._p = paragraph._element = None
#
# document = Document(r'source.docx')
# ps = [ paragraph for paragraph in document.paragraphs]
#
# delete_paragraph(ps[62])
#
# # ps[20].text = re.sub(r'{\$(.+?)}',respec,ps[20].text)
# # ps[20].style.font.name = 'Times New Roman'
# # ps[20].style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
# # ps[20].style.font.size = Pt(12)
# # document.paragraphs = document.paragraphs.insert(20,ps[20])
# # for p in ps:
# #     print(p)
# document.save('new.docx')



#initialization
app = Dispatch('Word.Application')
app.visible = True
# doc2 = app.Documents.Open(os.path.dirname(os.path.abspath(__file__)) + '/test.docx')
doc2 = None
doc1 = app.Documents.Open(os.path.dirname(os.path.abspath(__file__)) + '/test.docx')
doc1.ActiveWindow.View.ShowHiddenText = False

app2 = Dispatch('Excel.Application')
xbook = app2.Workbooks.Open(os.path.dirname(os.path.abspath(__file__)) + '/test.xlsm')

source=pd.read_excel(os.path.join(os.path.dirname(os.path.abspath(__file__)),'test.xlsm'),sheet_name=[0],header=0)
field_dict = collections.OrderedDict()
elements=source[0]
# elements = None
elements_dict = {}

data_list = [('合并资产负债表',(4,6),(110,32))]
label_list = [('注释',(7,4),(113,30))]
color_dict = {'橙色':49407,'蓝色':15773696,'绿色':5287936,'浅绿':5296274}

def replace_all(oldstr, newstr, regrex = False):
    app.Selection.Find.Execute(
        oldstr, False, False, regrex, False, False, True, 1, False, newstr, 2)

# app.Selection.Find.Execute('\{\$(*)\}', False, False, True, False, False, True, 1, False, '\\1', 2)
#Replace
def elements_pre():
    for i in range(len(elements)):
        elements_dict[elements['要素项目'][i]] = elements['内容'][i][2:-1]

# def elements_state1to3():
#     # with tqdm(total=len(elements_dict),bar_format='{l_bar}%s{bar}%s{r_bar}' % (Fore.BLUE, Fore.RESET)) as pbar:
#     with tqdm(total=len(elements_dict)) as pbar:
#         pbar.set_description('Replacing:')
#         for index,element in zip(range(len(elements_dict)),elements_dict):
#             newStr = elements_dict[element]
#             if not pd.isnull(newStr) and newStr!='-':
#                 replace_all(element,elements_dict[element])
#             pbar.update(1)
# def elements_state1to2():
#     # with tqdm(total=len(elements_dict),bar_format='{l_bar}%s{bar}%s{r_bar}' % (Fore.BLUE, Fore.RESET)) as pbar:
#     with tqdm(total=len(elements_dict)) as pbar:
#         pbar.set_description('Replacing:')
#         for index,element in zip(range(len(elements_dict)),elements_dict):
#             label = re.match(r'{\$(.+?)}',element).group(1)
#             newStr = '{' + 'M_{}_{}'.format(label,elements_dict[element]) + '}'
#             if not pd.isnull(newStr) and newStr!='-':
#                 replace_all(element,newStr)
#             pbar.update(1)
# def elements_state2to3():
#     replace_all('\{M_(*)_(*)\}','\\2',regrex=True)
# def elements_state2to1():
#     replace_all('\{M_(*)_(*)\}','{$\\1}',regrex=True)

def remove_color(color=65535):
    app.Selection.Find.Font.Color = color
    app.Selection.Find.Execute(FindText='', MatchCase=False, MatchWholeWord=False, MatchWildcards = False,
                               MatchSoundsLike = False, MatchAllWordForms = False, Forward = True,
                               Wrap = 1, Format = True, ReplaceWith='', Replace=2)

def set_find(ClearFormatting=True, Text="", ReplacementText="",Forward=True,
             Wrap = 1,Format = False,MatchCase = False,MatchWholeWord = False,
             MatchByte = True,MatchAllWordForms = False,MatchSoundsLike = False,MatchWildcards = True):
    if ClearFormatting:
        app.Selection.Find.ClearFormatting()
    app.Selection.Find.Text = Text
    app.Selection.Find.Replacement.Text = ReplacementText
    app.Selection.Find.Forward = Forward
    app.Selection.Find.Wrap = Wrap
    app.Selection.Find.Format = Format
    app.Selection.Find.MatchCase = MatchCase
    app.Selection.Find.MatchWholeWord = MatchWholeWord
    app.Selection.Find.MatchByte = MatchByte
    app.Selection.Find.MatchAllWordForms = MatchAllWordForms
    app.Selection.Find.MatchSoundsLike = MatchSoundsLike
    app.Selection.Find.MatchWildcards = MatchWildcards

def hide_change(hidden = False,name = None):
    app.Selection.Find.ClearFormatting()
    app.Selection.Find.Text = "\{T%s_*_T%s\}" % (name, name)
    app.Selection.Find.Replacement.Text = ""
    app.Selection.Find.Forward = True
    app.Selection.Find.Wrap = 1
    app.Selection.Find.Format = False
    app.Selection.Find.MatchCase = False
    app.Selection.Find.MatchWholeWord = False
    app.Selection.Find.MatchByte = True
    app.Selection.Find.MatchAllWordForms = False
    app.Selection.Find.MatchSoundsLike = False
    app.Selection.Find.MatchWildcards = True

    app.Selection.WholeStory()
    while app.Selection.Find.Execute():
        app.Selection.Font.Hidden = hidden

def all_com(now,rest):
    if rest != '':
        all_com(now, rest[1:])
        all_com(now+rest[0], rest[1:])
    else:
        if len(now) != 0:
            hide_change(hidden =True, name = ';'.join(now))

def hide_all_T(t_num):
    doc1.ActiveWindow.View.ShowHiddenText = True
    tl = [str(i) for i in range(1,t_num + 1)]
    all_com(now='',rest=''.join(tl))
    doc1.ActiveWindow.View.ShowHiddenText = False

def chosen_copy(now,rest,chosen):
    if rest != '':
        chosen_copy(now, rest[1:], chosen)
        chosen_copy(now+rest[0], rest[1:], chosen)
    else:
        if len(now) != 0 and chosen in now:
            name = ';'.join(now)
            replace_all("\{T%s_(*)_T%s\}" % (name, name),"{P_\\1_P}{T%s_\\1_T%s}" % (name, name),regrex=True)
            replace_all("^p^p", "^p")

def copy_chosen_T(chosen,t_num):
    tl = [str(i) for i in range(1,t_num + 1)]
    chosen_copy(now='', rest=''.join(tl), chosen=str(chosen))

def restore():
    doc1.ActiveWindow.View.ShowHiddenText = True
    app.Selection.WholeStory()
    replace_all('\{P_(*)_P\}', '', regrex=True)
    app.Selection.Font.Hidden = False
    doc1.ActiveWindow.View.ShowHiddenText = False

def replace_elements():
    for i,l in zip(data_list,label_list):
        row_num = i[2][0]-i[1][0]+1
        col_num = i[2][1]-i[1][1]+1
        with tqdm(total=row_num*col_num) as pbar:
            pbar.set_description(i[0])
            for row in range(row_num):
                for col in range(col_num):
                    data = xbook.sheets[i[0]].Cells(i[1][0]+row,i[1][1]+col).text.strip()
                    label = xbook.sheets[l[0]].Cells(l[1][0]+row,l[1][1]+col).text.strip()
                    if data!='' and data!='#DIV/0!' and data!='-':
                        replace_all(label,data)
                    pbar.update(1)

def remove_p():
    replace_all('\{P_(*)_P\}','\\1',regrex=True)

def remove_empty_row():
    doc1.ActiveWindow.View.ShowHiddenText = True
    for table in doc1.tables:
        if table.Cell(1,1).range.font.hidden == -1:
            continue
        row_index = 1
        for i in range(1, table.rows.count + 1):
            flag = True
            for j in range(1, table.columns.count + 1):
                try:
                    cell_str = table.Cell(row_index,j).range.text.split('\r')[0]
                    if cell_str != '':
                        flag = False
                        break
                except BaseException:
                    flag = False
                    break
            if flag:
                doc1.Range(table.Cell(row_index,1).range.start,table.Cell(row_index,table.columns.count).range.end).cells.Delete(2)
                row_index -= 1
            row_index += 1
    doc1.ActiveWindow.View.ShowHiddenText = False

def condition():
    app.Selection.Find.ClearFormatting()
    app.Selection.Find.Text = "\{T(*)_"
    app.Selection.Find.Replacement.Text = ""
    app.Selection.Find.Forward = True
    app.Selection.Find.Wrap = 1
    app.Selection.Find.Format = False
    app.Selection.Find.MatchCase = False
    app.Selection.Find.MatchWholeWord = False
    app.Selection.Find.MatchByte = True
    app.Selection.Find.MatchAllWordForms = False
    app.Selection.Find.MatchSoundsLike = False
    app.Selection.Find.MatchWildcards = True

    app.Selection.WholeStory()
    while app.Selection.Find.Execute():
        label_text = app.Selection.text[2:-1]
        print('ssss')

condition()
elements_pre()

hide_change(False,'')

# copy_chosen_T(chosen=2,t_num=4)
# hide_all_T(t_num=4)
# replace_elements()
# remove_color(65535)
# remove_empty_row()
restore()

#提交时
remove_p()










# re.match(r'(.+?)\n(.+)',a).group(2)
# xbook.sheets[5].Cells(20,6).Comment.Text()
# xbook.sheets[5].Range('F20').Comment.Text()
# doc1.ActiveWindow.View.ShowHiddenText = True
# replace_all('\{T33_*_T33\}^13', '', regrex=True)
#app.Selection.Range.ListFormat.ApplyListTemplate(ListTemplate=app.Selection.Range.ListFormat.ListTemplate, ContinuePreviousList=False)

def field_scan():
    for field in doc1.Fields:
        field.Select()
        if field.code.text.strip().isdigit():
            field_dict[field.code.text.strip()] = doc1.Range(0, field.application.selection.paragraphs[0].range.end).paragraphs.count
        if 'begin' in field.code.text.strip() or 'end' in field.code.text.strip():
            field_dict[field.code.text.strip()] = doc1.Range(0, field.application.selection.paragraphs[0].range.end).paragraphs.count
        if 'tbegin' in field.code.text.strip() or 'tend' in field.code.text.strip():
            field_dict[field.code.text.strip()] = doc1.Range(0, field.application.selection.paragraphs[0].range.end).paragraphs.count
        print(field.code.text.strip())

def update_field_dict(begin_name, end_name, enum=0 , pnum=0):

    flag = False
    diff = 0
    if enum == 0:
        diff = field_dict[end_name] - field_dict[begin_name] + 1
    elif enum == 1:
        diff = pnum
    for key in list(field_dict):
        if not flag:
            if key == begin_name or key == end_name:
                if enum == 0:
                    field_dict.pop(key)
                elif enum == 1:
                    flag = True
                    continue
            if key == end_name:
                flag = True
                continue
        if flag:
            field_dict[key] -= diff

def field_trans():
    for field in doc2.Fields:
        locs = field.code.text.strip().split('#')
        if len(locs) == 2:
            field.Select()
            i = doc2.Range(0,field.application.selection.paragraphs[0].range.end).paragraphs.count
            # field.Delete()
            doc1.Range(doc1.paragraphs[field_dict[locs[0]]].range.start, doc1.paragraphs[field_dict[locs[1]] - 2].range.end).Copy()
            doc2.paragraphs[i-1].range.PasteAndFormat(16)

def pop_field(block_name):
    begin_name = '{}begin'.format(block_name)
    end_name = '{}end'.format(block_name)
    rangeStart = doc1.paragraphs[field_dict[begin_name]-1].range.start
    rangeEnd = doc1.paragraphs[field_dict[end_name] - 1].range.end
    doc1.Range(rangeStart, rangeEnd).Delete()

    update_field_dict(begin_name, end_name)

def table_add_col(table_name, num=1):
    table = field2table(table_name)
    begin_name = '{}tbegin'.format(table_name)
    end_name = '{}tend'.format(table_name)
    for i in range(num):
        table.columns.Add()
        table.columns.last.Cells.verticalalignment = 1
        update_field_dict(begin_name, end_name, 1, -table.rows.count)
    table.AutoFitBehavior(2)

def table_del_col(table_name, num=1):
    table = field2table(table_name)
    begin_name = '{}tbegin'.format(table_name)
    end_name = '{}tend'.format(table_name)
    for i in range(num):
        table.columns.last.Delete()
        update_field_dict(begin_name, end_name, 1, table.rows.count)
    table.AutoFitBehavior(2)

def remove_zero_row(block_name, row_off = 4,col_off = 2):
    begin_name = '{}tbegin'.format(block_name)
    end_name = '{}tend'.format(block_name)

    table = field2table(block_name)

    row_index = row_off
    for i in range(row_off, table.rows.count + 1):
        flag = True
        for j in range(col_off, table.columns.count + 1):
            cell_str = table.Cell(row_index,j).range.text.split('\r')[0]
            if not cell_str.isdigit() or float(cell_str) != 0:
                flag = False
                break
        if flag:
            doc1.Range(table.Cell(row_index,1).range.start,table.Cell(row_index,table.columns.count).range.end).cells.Delete(2)
            update_field_dict(begin_name, end_name, 1, table.columns.count + 1)
            row_index -= 1
        row_index += 1

def field2table(table_name):
    begin_name = '{}tbegin'.format(table_name)
    end_name = '{}tend'.format(table_name)
    rangeStart = doc1.paragraphs[field_dict[begin_name] - 1].range.end
    rangeEnd = doc1.paragraphs[field_dict[end_name] - 1].range.start
    table = doc1.Range(rangeStart, rangeEnd).tables[0]
    return table

#----------------------------------------Testing--------------------------------------
field_scan()
pop_field('有担保')
pop_field('无担保')
pop_field('标号测试')
pop_field('债券发行')

# Table Options

doc1.paragraphs[25].range.Copy()
doc1.paragraphs[27].range.PasteAndFormat(16)

# a = doc1.paragraphs[24].range.highlightcolorindex
# b = doc1.paragraphs[22].range.font.colorindex

table_add_col('资产负债', 2)
table_del_col('资产负债', 1)

pop_field('货币资金1')
pop_field('货币资金2')
pop_field('货币资金3')
pop_field('货币资金4')

field_trans()

#insert col
for i in range(1, doc1.tables[0].rows.count + 1):
    doc1.tables[0].Cell(i, doc1.tables[0].columns.count).range.text = 'content_{}'.format(i)

doc1.tables[0].Cell(1,1).range.text = '具体项目'

#remove zero row
remove_zero_row('资产结构')

doc1.SaveAs(os.path.join(os.path.dirname(os.path.abspath(__file__)),'result.docx'))
doc1.Close(0)



print('Done!')

